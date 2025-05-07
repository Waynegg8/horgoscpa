#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文檔內容提取工具 - 為多個模組提供Word文檔讀取功能
作者: Claude
日期: 2025-05-07
"""

import os
import sys
import logging
import importlib
import subprocess
from typing import Dict, Any, Optional

# 設置日誌
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("word_extractor")

def ensure_dependencies():
    """確保依賴庫已安裝"""
    required_packages = [
        "mammoth", "python-docx", "bs4", "lxml", "docx2txt"
    ]
    
    for package in required_packages:
        try:
            importlib.import_module(package)
        except ImportError:
            logger.info(f"安裝缺失的依賴: {package}")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

def extract_word_content_direct(docx_path: str) -> str:
    """直接提取Word文檔內容為純文本"""
    logger.info(f"使用直接提取法從Word獲取內容: {docx_path}")
    try:
        import docx
        
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                full_text.append(" ".join(row_text))
        
        logger.info(f"直接提取了 {len(full_text)} 個段落")
        return "\n\n".join(full_text)
    except Exception as e:
        logger.error(f"直接提取Word內容時出錯: {str(e)}")
        return ""

def extract_word_with_mammoth(docx_path: str) -> str:
    """使用mammoth提取Word文檔為HTML，然後提取純文本"""
    logger.info(f"使用mammoth從Word獲取HTML: {docx_path}")
    try:
        import mammoth
        from bs4 import BeautifulSoup
        
        with open(docx_path, 'rb') as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
            
            # 使用BeautifulSoup提取純文本
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text()
            
            logger.info(f"Mammoth提取的文本長度: {len(text_content)}")
            return text_content
    except Exception as e:
        logger.error(f"Mammoth提取Word內容時出錯: {str(e)}")
        return ""

def extract_word_with_docx2txt(docx_path: str) -> str:
    """使用docx2txt提取Word文檔為純文本"""
    logger.info(f"使用docx2txt從Word獲取純文本: {docx_path}")
    try:
        import docx2txt
        
        text = docx2txt.process(docx_path)
        logger.info(f"Docx2txt提取的文本長度: {len(text)}")
        return text
    except Exception as e:
        logger.error(f"Docx2txt提取Word內容時出錯: {str(e)}")
        return ""

def extract_word_document_content(docx_path: str) -> Optional[str]:
    """整合多種方法提取Word文檔內容，返回最佳結果"""
    logger.info(f"開始提取Word文檔內容: {docx_path}")
    
    # 確保依賴庫已安裝
    ensure_dependencies()
    
    # 使用多種方法提取內容
    results = {
        'direct_text': extract_word_content_direct(docx_path),
        'mammoth_text': extract_word_with_mammoth(docx_path),
        'docx2txt': extract_word_with_docx2txt(docx_path)
    }
    
    # 選擇最佳結果（內容最多的）
    content_lengths = {
        'direct_text': len(results['direct_text']),
        'mammoth_text': len(results['mammoth_text']),
        'docx2txt': len(results['docx2txt'])
    }
    
    best_method = max(content_lengths.items(), key=lambda x: x[1])[0]
    best_content = results[best_method]
    
    logger.info(f"選擇的最佳提取方法: {best_method}, 長度: {content_lengths[best_method]}")
    
    if not best_content:
        logger.warning("無法提取Word文檔內容")
        return None
    
    return best_content

if __name__ == "__main__":
    """如果直接執行此腳本，將測試一個Word文檔的提取"""
    if len(sys.argv) < 2:
        print("用法: python word_extractor.py <Word文檔路徑>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    content = extract_word_document_content(docx_path)
    
    if content:
        print(f"成功提取Word文檔內容, 長度: {len(content)}")
        print("內容前200個字符:")
        print(content[:200] + "...")
    else:
        print("提取失敗")